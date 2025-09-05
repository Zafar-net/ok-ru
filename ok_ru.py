from customtkinter import *
import time
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.common.exceptions import NoSuchElementException
from docx import Document
from docx.shared import Inches
import re  # Linklarni filtrlash uchun

set_appearance_mode("dark")
set_default_color_theme("dark-blue")

root = CTk()
root.title('Odonoklassniki_Strike')
root.iconbitmap('ok_ru.ico')
root.geometry("600x300")

t, t1, t2, t3, t4, t5, t6, t7 = "", "", "", "", "", "", "", "",


def selenium1(value):
    global d
    d = value
    if d == "В видеоролике(-ах)" or d == "В теме(-ах) в группе":
        global t, t1, t2, t3, t4, t5, t6, t7
        t = "...."
        t1 = "Экстремизм/терроризм"
        t2 = "Разжигание межнациональной розни"
        t3 = "Информация от нежелательных организаций"
        t4 = "Массовые беспорядки/мероприятия"
        t5 = "Сцены насилия и жестокий контент"
        t6 = "Недостоверная информация"
        t7 = "Спамерский контент"

        box1.configure(values=[t, t1, t2, t3, t4, t5, t6, t7])
        box1.set(t)
        box1.pack(pady=10)
    print(f"Tanlangan kategoriya {value}")


def box1(value):
    global b1
    b1 = value
    if b1 == "Экстремизм/терроризм":
        b1 = "//*[@value='EXTREMISM_TERRORISM']"

    if b1 == "Разжигание межнациональной розни":
        b1 = "//*[@value='ETHNIC']"

    if b1 == "Информация от нежелательных организаций":
        b1 = "//*[@value='INFORMATION_FROM_UNWANTED_ORGANIZATIONS']"

    if b1 == "Массовые беспорядки/мероприятия":
        b1 = "//*[@value='MASS_RIOTS']"

    if b1 == "Сцены насилия и жестокий контент":
        b1 = "//*[@value='VIOLENCE']"

    if b1 == "Недостоверная информация":
        b1 = "//*[@value='INACCURATE_INFORMATION']"

    if b1 == "Спамерский контент":
        b1 = "//*[@value='SPAM_CONTENT']"

    print(f" Tanlangan kategoriya 2 {value}")

def legenda():
    if d == "В видеоролике(-ах)":
        h = "//*[@value='IN_VIDEO']"
        f = "links\\link.txt"
        c = b1
        if b1 == "//*[@value='ETHNIC']":
            c2 = "Данное высказывание нарушает положения статьи 282 Уголовного кодекса РФ"
            t = "Данное высказывание нарушает положения статьи 282 Уголовного кодекса РФ(возбуждение ненависти либо " \
                "вражды, а равно унижение человеческого достоинства).Прошу: 1.  Провести проверку по данному факту. " \
                "2.  Принять соответствующие меры реагирования в рамках действующего законодательства. 3.  При " \
                "необходимости, привлечь виновных лиц к административной или уголовной ответственности. 4.  При " \
                "возможности, удалить противоправный контент. "

        else:
            # c2 = "Данное высказывание нарушает положения статьи 282 Уголовного кодекса РФ"
            # t = "Данное высказывание нарушает положения статьи 282 Уголовного кодекса РФ(возбуждение ненависти либо " \
            #     "вражды, а равно унижение человеческого достоинства).Прошу: 1.  Провести проверку по данному факту. " \
            #     "2.  Принять соответствующие меры реагирования в рамках действующего законодательства. 3.  При " \
            #     "необходимости, привлечь виновных лиц к административной или уголовной ответственности. 4.  При " \
            #     "возможности, удалить противоправный контент. "
            t = "УК РФ Статья 205.2. Публичные призывы к осуществлению террористической деятельности, публичное " \
                "оправдание терроризма или пропаганда терроризма Призывы к терроризму (ч. 2 ст. 205.2 УК РФ) "
            c2 = "Пропаганда и заблуждение мирного народов мира"

        s = "images\\Video.png"
        strike(h, f, c, c2, t, s)

    if d == "В имени":
        h = "//*[@value='IN_NAME']"
        f = "links\\profiles.txt"
        c = "//*[@value='ETHNIC']"
        c2 = "Призывы к терроризму (ч. 2 ст. 205.2 УК РФ)"
        t = ""
        s = "images\\profiles.png"
        strike(h, f, c, c2, t, s)

    if d == "В названии группы":
        h = "//*[@value='IN_GROUP_NAME']"
        f = "links\\gr.txt"
        c = ""
        c2 = "Публичная страница распространяет запрещенную в РФ информацию:(ст. 207.3 УК РФ)"
        t = ""
        s = "images\\gr.png"
        strike(h, f, c, c2, t, s)

    if d == "В заметке(-ах) пользователя":
        h = "//*[@value='IN_NOTES']"
        f = "links\\notes.txt"
        c = "//*[@value='EXTREMISM_TERRORISM']"
        c2 = "Призывы к терроризму (ч. 2 ст. 205.2 УК РФ)"
        t = ""
        s = "images\\notes.png"
        strike(h, f, c, c2, t, s)

    if d == "На фото/изображении(-ях)":
        h = "//*[@value='IN_PHOTO']"
        f = "links\\photo.txt"
        c = "//*[@value='EXTREMISM_TERRORISM']"
        c2 = "Призывы к терроризму (ч. 2 ст. 205.2 УК РФ)"
        t = ""
        s = "images\\photo.png"
        strike(h, f, c, c2, t, s)

    if d == "В теме(-ах) в группе":
        h = "//*[@value='IN_GROUP_THEME']"
        f = "links\\topic.txt"
        c = b1
        # Указывается информация, нарушающая законодательство РФ (конкретные нарушения)
        # Распространение недостоверной общественно значимой информации РФ (ст. 13.15 КоАП РФ )
        #Публикация содержит заведомо ложную информацию, распространяемую под видом достоверных сообщений, что создает угрозу общественной безопасности (ст. 13.15 КоАП РФ - Злоупотребление свободой массовой информации).
        c2 = "Указывается информация, нарушающая законодательство РФ (ст.13.15 КоАП РФ)"
        t = ""
        s = "images\\2.png"
        strike(h, f, c, c2, t, s)
    else:
        print("")

def data():
    v_url =entry.get()
    count = 1
    #C:\Users\User\AppData\Local\Google\Chrome\User Data
    profile = [27, 29]
    for x in profile:
        options = Options()
        options.add_argument(
            r"--user-data-dir=D:\\Programma\\folder_interview\\Selenium_Auto\\tkinter\\User Data")  # e.g. C:\Users\You\AppData\Local\Google\Chrome\User Data widget_w __vis
        options.add_argument(fr'--profile-directory=Profile {x}')  # e.g. Profile 3
        options.add_argument("--disable-extensions")
        options.add_experimental_option("useAutomationExtension", False)
        options.add_argument('--no-sandbox')
        service = Service(executable_path='chromedriver.exe')
        driver = webdriver.Chrome(service=service, options=options)

        driver.get(v_url)
        time.sleep(2)
        # driver.find_element(By.XPATH, '(//*[@class="video-card js-movie-card "])[1]').click() video-card js-movie-card  __responsive
        # video-card js-movie-card group-video-card __responsive
        #

        urls = []
        imgs = []
        data = []

        try:
            element = driver.find_element(By.XPATH, f'(//*[@class="video-card js-movie-card "])[{count}]')
        except NoSuchElementException:
            element = driver.find_element(By.XPATH, f'(//*[@class="video-card js-movie-card  __responsive"])[{count}]')

        if element:
            a = True
        else:
            a = False
            break
        while a:
            time.sleep(2)
            try:
                time.sleep(2)
                if element:
                    time.sleep(2)
                    driver.execute_script(f"window.scrollBy(0, {80 * count})")
                    time.sleep(2)
                    try:
                        driver.find_element(By.XPATH,
                                            f'(//*[@class="video-card js-movie-card  __responsive"])[{count}]').click()
                        time.sleep(4)
                    except:
                        driver.find_element(By.XPATH,
                                            f'(//*[@class="video-card js-movie-card  __responsive"])[{count}]').click()
                        time.sleep(4)

                    a = True
                else:
                    a = False
                    # break
            except NoSuchElementException:
                break
            time.sleep(4)
            driver.get_screenshot_as_file(f"images\\{count}.png")
            p_image = f"images\\{count}.png"
            imgs.append(p_image)
            time.sleep(2)
            count += 1
            current_url = driver.current_url
            urls.append(current_url)
            time.sleep(2)
            print(f"{urls} \n {imgs}")
            driver.get(v_url)

        # Rasm va linklarni bitta o'zgaruvchiga yig'ish
        for url, img in zip(urls, imgs):
            data.append({"url": url, "img": img})
        print(data)
        #
        doc = Document()
        doc.add_heading("Obyekt ichidagi videolarning havolasi hamda boshlang'ich rasmi", level=1)
        table = doc.add_table(rows=1, cols=3)  # 2 ustunli jadval
        table.style = 'Table Grid'

        # Birinchi qatorga sarlavha qo'shish
        header_cells = table.rows[0].cells
        header_cells[0].text = "T/r"
        header_cells[1].text = "Link"
        header_cells[2].text = "Rasm"

        # Ma'lumotlarni jadvalga qo'shish
        c = 0
        for item in data:
            c += 1
            row_cells = table.add_row().cells

            row_cells[0].text = str(c)

            # Tekstni birinchi ustunga qo'shish
            row_cells[1].text = item["url"]

            # Rasmni ikkinchi ustunga qo'shish
            paragraph = row_cells[2].paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(item["img"], width=Inches(2))

        # Hujjatni saqlash
        doc.save(f"{count - 1}.docx")
        # Word hujjatini ochish
        doc = Document(f"{count - 1}.docx")

        # Tablitsadan topilgan barcha linklarni yig'ish uchun ro'yxat
        links = []

        # Tablitsalardan ma'lumotlarni o'qish
        for table in doc.tables:  # Hujjatdagi barcha tablitsalarni tekshirish
            for row in table.rows:  # Har bir qatorni tekshirish
                for cell in row.cells:  # Har bir hujayra
                    for paragraph in cell.paragraphs:  # Hujayradagi barcha paragraphlarni tekshirish
                        text = paragraph.text.strip()  # Matnni olish
                        # Linkni topish uchun filtr: `https://` bilan boshlanadigan qism
                        found_links = re.findall(r'https?://[^\s]+', text)  # Linklarni ajratib olish
                        links.extend(found_links)

        # Natijalarni chiqarish
        print("Tablitsadan topilgan linklar:")
        for link in links:
            print(link)

        # Linklarni link.txt fayliga saqlash
        with open("links\\link.txt", "w", encoding="utf-8") as file:
            for link in links:
                file.write(link + "\n")  # Har bir link yangi qatorga yoziladi

        print(f"\n{len(links)} ta link 'link.txt' fayliga muvaffaqiyatli saqlandi!")

        driver.quit()


def strike(h, f, c, c2, t, s):

    profile = [12, 13, 14, 15, 16, 17, 18, 19, 20, 21, 22, 23, 24, 25, 26, 27, 30, 34, 35, 36, 37, 42, 44, 46, 48, 49, 50, 51,]
    for p in profile:
        options = Options()
        options.add_argument(
            r"--user-data-dir=D:\\Programma\\folder_interview\\Selenium_Auto\\tkinter\\User Data")  # e.g. C:\Users\You\AppData\Local\Google\Chrome\User Data
        options.add_argument(fr'--profile-directory=Profile {p}')  # e.g. Profile 3 67 48
        options.add_argument("--disable-extensions")
        options.add_experimental_option("useAutomationExtension", False)
        options.add_argument('--no-sandbox')
        service = Service(executable_path='chromedriver.exe')
        driver = webdriver.Chrome(service=service, options=options)
        # link.txt faylni ochib ichidagilarni list1 ga o'zlashtirib oladi
        file = open(f, "r")
        data = file.read()
        list1 = data.replace('\n', ' ')
        file.close()

        # string tipidagi o'zgaruvchi ichidagi ma'lumotlarni probel bilan uzilgan joyidan listga qo'shib oladi
        def Convert(string):
            li = list(string.split(" "))
            return li

        link1 = Convert(list1)

        # Ariza berish boshlanishi
        driver.get("https://ok.ru/help/ask/COMPLAINT")
        time.sleep(2)
        driver.find_element(By.XPATH, h).click()
        time.sleep(3)
        if d != "В названии группы":
            driver.find_element(By.XPATH, c).click()
        # driver.find_element(By.XPATH, c).click()
        time.sleep(1)

        driver.find_element(By.XPATH, "//*[@id='field_forbiddenInformation']").send_keys(c2)

        if d == "В видеоролике(-ах)":
            driver.find_element(By.XPATH, "//*[@class='itx js-textarea']").send_keys(t)
        else:
            time.sleep(1)
        time.sleep(1)

        # ariza beriladigan linklarni joylash
        count = 0
        for x in link1:
            count += 1
            if count == 1:
                driver.find_element(By.XPATH, "//*[@class='it js-lastline']").send_keys(x)
                driver.find_element(By.XPATH, "//*[@class='it_w-add al js-multiline-append']").click()
            else:
                driver.find_element(By.XPATH, "//*[@class='it js-multiline-input js-lastline']").send_keys(x)
                if x == '':
                    continue
                else:
                    if count != 10:
                        driver.find_element(By.XPATH, "//*[@class='it_w-add al js-multiline-append']").click()
                    else:
                        break

        driver.get_screenshot_as_file(f"images\\{x}.png")
        time.sleep(2)
        driver.find_element(By.XPATH, "//*[@class='form-actions_yes button-pro js-submit']").click()
        time.sleep(2)
        print(f'{p} profil tugatdi')
        driver.quit()

    # Profillarga bo'lib havola berish
    # profile = [12, 13, 14, 15]
    #
    # # Havolalarni o'qib olish
    # with open(f, "r") as file:
    #     data = file.read()
    #     list1 = data.replace('\n', ' ')
    #
    # def Convert(string):
    #     return list(filter(None, string.split(" ")))
    #
    # all_links = Convert(list1)
    # chunk_size = 10  # Har bir arizada yuboriladigan havolalar soni
    #
    # # Havolalarni 10 tadan guruhlarga ajratamiz
    # link_chunks = [all_links[i:i + chunk_size] for i in range(0, len(all_links), chunk_size)]
    #
    # for chunk_index, chunk in enumerate(link_chunks):
    #     # Har bir 10 talik guruh uchun navbatdagi profilni tanlaymiz
    #     p = profile[chunk_index % len(profile)]
    #     print(f"\n{chunk_index + 1}-guruh (Profil {p}): {len(chunk)} ta havola")
    #
    #     options = Options()
    #     options.add_argument(
    #         r"--user-data-dir=D:\\Programma\\folder_interview\\Selenium_Auto\\tkinter\\User Data")
    #     options.add_argument(fr'--profile-directory=Profile {p}')
    #     options.add_argument("--disable-extensions")
    #     options.add_experimental_option("useAutomationExtension", False)
    #     options.add_argument('--no-sandbox')
    #     service = Service(executable_path='chromedriver.exe')
    #     driver = webdriver.Chrome(service=service, options=options)
    #
    #     try:
    #         # Ariza berish jarayoni
    #         driver.get("https://ok.ru/help/ask/COMPLAINT")
    #         time.sleep(2)
    #         driver.find_element(By.XPATH, h).click()
    #         time.sleep(3)
    #         driver.find_element(By.XPATH, c).click()
    #         time.sleep(1)
    #
    #         driver.find_element(By.XPATH, "//*[@id='field_forbiddenInformation']").send_keys(c2)
    #
    #         if d == "В видеоролике(-ах)":
    #             driver.find_element(By.XPATH, "//*[@class='itx js-textarea']").send_keys(t)
    #         else:
    #             time.sleep(1)
    #         time.sleep(1)
    #
    #         # Havolalarni kiritish
    #         for i, link in enumerate(chunk):
    #             if i == 0:
    #                 driver.find_element(By.XPATH, "//*[@class='it js-lastline']").send_keys(link)
    #                 if len(chunk) > 1:
    #                     driver.find_element(By.XPATH, "//*[@class='it_w-add al js-multiline-append']").click()
    #             else:
    #                 driver.find_element(By.XPATH, "//*[@class='it js-multiline-input js-lastline']").send_keys(link)
    #                 if i != len(chunk) - 1:
    #                     driver.find_element(By.XPATH, "//*[@class='it_w-add al js-multiline-append']").click()
    #
    #         # Screenshot va yuborish
    #         driver.get_screenshot_as_file(f"{s}_profil{p}_chunk{chunk_index + 1}.png")
    #         time.sleep(2)
    #         driver.find_element(By.XPATH, "//*[@class='form-actions_yes button-pro js-submit']").click()
    #         time.sleep(2)
    #
    #         print(f"Profil {p} dan {len(chunk)} ta havolaga ariza yuborildi!")
    #
    #     except Exception as e:
    #         print(f"Xatolik yuz berdi (Profil {p}): {str(e)}")
    #
    #     finally:
    #         driver.quit()
    #         time.sleep(1)
    #
    # print("\nBarcha havolalar profillar orqali muvaffaqiyatli taqsimlab yuborildi!")


# TabView yaratish
tabview = CTkTabview(master=root)
tabview.pack(pady=10, padx=10, fill="both", expand=True)

# 2 ta tab qo'shish
tab1 = tabview.add("Videoga spam berish")
tab2 = tabview.add("Data")

#Birinchi tab uchun content (ComboBoxlar)
frame1 = CTkFrame(master=tab1)
frame1.pack(pady=20, padx=20, fill="both", expand=True)
box = CTkComboBox(master=frame1, values=["....", "В видеоролике(-ах)", "В заметке(-ах) пользователя", "В имени", "В названии группы",
                                       "На фото/изображении(-ях)", "В теме(-ах) в группе"], width=300, fg_color="coral",
                  border_color="#FBAB7E", dropdown_fg_color="coral", command=selenium1)
box.pack(pady=10)
box1 = CTkComboBox(master=frame1, values=[t, t1, t2, t3, t4, t5, t6, t7], width=300,
                      fg_color="coral",
                      border_color="#FBAB7E", dropdown_fg_color="coral", command=box1, text_color="white",
                      dropdown_text_color="white")
box1.pack(pady=10)
btn = CTkButton(master=frame1, text="Ishga tushirish", command=legenda)
btn.pack(pady=20)

frame2 = CTkFrame(master=tab2)
frame2.pack(pady=20, padx=20, fill="both", expand=True)

entry = CTkEntry(master=frame2,
                width=300,
                placeholder_text="Havolani shu yerga kiriting...",
                fg_color="coral",
                border_color="#FBAB7E",
                text_color="white")
entry.pack(pady=20)

tn = CTkButton(master=frame2, text="Data", command=data)
tn.pack(pady=10)
root.mainloop()